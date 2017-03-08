from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from mpl_toolkits.axes_grid.parasite_axes import SubplotHost
from mpl_toolkits.axes_grid1 import host_subplot
import mpl_toolkits.axisartist as AA
import seaborn as sns
import datetime

#
# Use Document to make a word doc
# with sessions so that sessions can 
# be viewed for discussion with engineering team
#
document = Document()
date_ = str(datetime.date.today())
date_.replace(" ","_")

document.add_heading('Classify Biotrak Sessions from Trial', 0)

df_ratings = pd.read_csv('Trial_sessions_report.csv')
df_sessions = pd.read_json('sessions_latest.txt')

document.add_heading('Start of trial user sessions', level=1)

#
# Take out all except trial users
#

df_trial_users = pd.read_csv('TrialUsers.csv',header=None)
df_sessions = df_sessions[df_sessions['userid'].isin(df_trial_users[0])]
df_sessions['session_num'] = df_sessions.index
df_sessions.reset_index(inplace=True)
df_sessions['Not_effective'] = df_ratings['Not_effective']

document.add_page_break()
s6 = df_sessions.copy()
#
# Loop through all of the trial user sessions to 
# create a graph that can be evaluated
#
for i in range(0, len(s6['session_data'])):
    plt.clf()
    y_avg_0 = map(int,filter(None,s6['session_data'].iloc[i]['average_data'].replace(" ","").split(",")))
    y_raw_0 = map(int,filter(None,s6['session_data'].iloc[i]['raw_data'].replace(" ","").split(",")))
    
    ## Invert signal per Scott's instructions
    y_avg_0 = [ 4095 - x for x in y_avg_0]
    
    y_display_0 = [0.0,10.0]
    if ('graphed_data' in s6['session_data'].iloc[i]):
        if type(s6['session_data'].iloc[i]['graphed_data'])==unicode:
            y_display_0 = map(float,filter(None,s6['session_data'].iloc[i]['graphed_data'].replace(" ","").split(",")))
        else:
            y_display_0 = map(float,filter(None,s6['session_data'].iloc[i]['graphed_data']))
    # App has put in incorrectly formatted last entry -- just discard it
    y_avg_0 = y_avg_0[:-1]
    y_raw_0 = y_raw_0[:-1]
    
    text = ""
    for key,value in  (s6['Session_data_header'].iloc[i]).iteritems():    
         text += key + ": " +value+", "
    text = text[:-2]
    ax = host_subplot(111, axes_class=AA.Axes)
    ax.set_title('Session # {0} data for user: {1} -- Session Type: {2}\n Name: {3}  | Time: {4}\n Session Header info: {5}\n'.format(i,
            s6['userid'].iloc[i],s6['Session_type'].iloc[i],s6['Session_name'].iloc[i],
             str(s6['Local_date'].iloc[i])+" : "+str(s6['Start_local_time'].iloc[i]),
            text))
    x2_offset = 0
    x3_offset = 0
    if ( len(y_avg_0) > len(y_display_0)):
        x3_offset = len(y_avg_0) - len(y_display_0)
    if ( len(y_avg_0) > len(y_raw_0)):
        x2_offset = len(y_avg_0) - len(y_raw_0)
        
    
    # X label and left y label
    ax.set_xlabel('time sec')
    ax.set_ylabel('avg_data')
    
    # X values for each plot
    x1 = np.arange(0,len(y_avg_0))/60.0
    x2 = np.arange(x2_offset,x2_offset+len(y_raw_0))/60.0
    x3 = np.arange(x3_offset,x3_offset+len(y_display_0))/60.0
    
    # Second y axis -- right
    ax2 = ax.twinx()
    
    # Third y axis -- ofset to the right
    offset=60
    ax3 = ax.twinx()   
    new_fixed_axis = ax3.get_grid_helper().new_fixed_axis
    ax3.axis["right"] = new_fixed_axis(loc="right",
                                        axes=ax3,
                                        offset=(offset, 0))
    
    
    if len(y_avg_0)>1:
        ax.set_ylim(0, 4095 * 1.1)  
    if len(y_raw_0)>1:
        ax2.set_ylim(0, 4095 * 1.1) 
    if len(y_display_0)>1:
        ax3.set_ylim(0, 10.5)
        ax3.fill_between(x3, y_display_0)
    
    # Set out plotted data
    l1 = ax.plot(x1, y_avg_0, 'b', label='average_data',alpha=0.5)
    l2 = ax2.plot(x2, y_raw_0, 'r', label='raw_data', alpha=0.5, )
    l3 = ax3.plot(x3, y_display_0, '#425e74', label='User Session', alpha=0.5, )
        
    # Adust Y axis parameters
    ax.axis["left"].label.set_color(l1[0].get_color())
    ax.tick_params(axis='y', color=l1[0].get_color(), labelcolor=l1[0].get_color())
    ax.yaxis.label.set_color(l1[0].get_color())
   
    ax.ticklabel_format(color=l1[0].get_color())
    ax.set_xlim(0,600) 
    
    # Adjust second y axis parameters
    ax2.axis["right"].label.set_color(l2[0].get_color())
    ax2.tick_params(axis='y', color=l2[0].get_color(), labelcolor=l2[0].get_color())
    ax2.yaxis.label.set_color(l1[0].get_color())
    ax2.set_ylabel('raw_data', color=l2[0].get_color())       

    # Adjust third y axis parameters
    ax3.tick_params(axis='y', color=l3[0].get_color())
    ax3.yaxis.label.set_color(l3[0].get_color())
    ax3.set_ylabel('User Session aka "graph_data"', color=l3[0].get_color())

    lns = l1+l2+l3
    labs = [l.get_label() for l in lns]
    
    ax.legend(loc='best', fancybox=True, framealpha=0.5)
    # 
    # Save each figure to a file 
    #
    plt.savefig('btrak_fig.png')
    #
    # Load it into the word document before continuing
    document.add_picture('btrak_fig.png', width=Inches(5))
    #
    # Add information to be used with the graph
    #
    document.add_paragraph(
        'Session {0} --Problem Session? {1}'.format(s6['session_num'].iloc[i], s6['Not_effective'].iloc[i]))
       
    document.add_page_break()
#
# Save the document
#
document.save('btrak_trial_classify.docx')
